# utils/docx_reader.py
import sys
import argparse
from pathlib import Path
from docx import Document

WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
RNS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
DNS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

HEADING_MAP = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
}


def _heading_level(paragraph):
    style = paragraph.style
    if style and style.name in HEADING_MAP:
        return HEADING_MAP[style.name]
    try:
        if style and style.paragraph_format:
            lvl = style.paragraph_format.outline_level
            if lvl is not None and 0 <= lvl <= 5:
                return lvl + 1
    except AttributeError:
        pass
    return None


def _is_list_item(paragraph):
    return paragraph._element.find(f".//{WNS}numPr") is not None


def _is_ordered_list(paragraph):
    """Check numbering part to determine if list is ordered (numbered) or bullet."""
    try:
        numPr = paragraph._element.find(f".//{WNS}numPr")
        if numPr is None:
            return False
        numId_el = numPr.find(f"{WNS}numId")
        ilvl_el = numPr.find(f"{WNS}ilvl")
        if numId_el is None:
            return False
        num_id = numId_el.get(f"{WNS}val")
        ilvl = ilvl_el.get(f"{WNS}val", "0") if ilvl_el is not None else "0"

        part = paragraph.part
        if not hasattr(part, "numbering_part") or part.numbering_part is None:
            return False
        numbering = part.numbering_part._element

        # Resolve numId → abstractNumId
        abstract_id = None
        for num in numbering.findall(f"{WNS}num"):
            if num.get(f"{WNS}numId") == num_id:
                ref = num.find(f"{WNS}abstractNumId")
                if ref is not None:
                    abstract_id = ref.get(f"{WNS}val")
                break
        if abstract_id is None:
            return False

        # Find numFmt for the matching level
        for abstract_num in numbering.findall(f"{WNS}abstractNum"):
            if abstract_num.get(f"{WNS}abstractNumId") == abstract_id:
                for lvl in abstract_num.findall(f"{WNS}lvl"):
                    if lvl.get(f"{WNS}ilvl") == ilvl:
                        num_fmt = lvl.find(f"{WNS}numFmt")
                        if num_fmt is not None:
                            return num_fmt.get(f"{WNS}val", "bullet") != "bullet"
    except Exception:
        pass
    # Fallback: check style name
    style_name = paragraph.style.name if paragraph.style else ""
    return "List Number" in style_name


def _rpr_fmt(rPr):
    """Return formatting flags from a w:rPr element."""
    if rPr is None:
        return {}

    def active(tag):
        el = rPr.find(f"{WNS}{tag}")
        if el is None:
            return False
        val = el.get(f"{WNS}val", "true").lower()
        return val not in ("false", "0", "off")

    u_el = rPr.find(f"{WNS}u")
    underline = u_el is not None and u_el.get(f"{WNS}val", "single") not in ("none", "")

    # Font color: w:color w:val="RRGGBB" (ignore "auto")
    color = None
    color_el = rPr.find(f"{WNS}color")
    if color_el is not None:
        val = color_el.get(f"{WNS}val", "auto")
        if val.lower() not in ("auto", ""):
            color = f"#{val.upper()}"

    return {
        "bold": active("b"),
        "italic": active("i"),
        "strike": active("strike"),
        "underline": underline,
        "color": color,
    }


def _apply_fmt(text, fmt):
    if not text:
        return text
    if fmt.get("strike"):
        text = f"~~{text}~~"
    if fmt.get("bold") and fmt.get("italic"):
        text = f"***{text}***"
    elif fmt.get("bold"):
        text = f"**{text}**"
    elif fmt.get("italic"):
        text = f"*{text}*"
    # Underline has no standard markdown; use HTML
    if fmt.get("underline") and not fmt.get("strike"):
        text = f"<u>{text}</u>"
    # Font color via HTML span
    if fmt.get("color"):
        text = f'<span style="color:{fmt["color"]}">{text}</span>'
    return text


def _run_text(r_elem):
    """Get text from a w:r element, preserving preserved spaces."""
    t = r_elem.find(f"{WNS}t")
    return t.text or "" if t is not None else ""


def _para_to_md_text(para_elem, doc):
    """Walk paragraph children to build inline-formatted markdown text."""
    parts = []

    for child in para_elem:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "r":
            fmt = _rpr_fmt(child.find(f"{WNS}rPr"))
            text = _run_text(child)
            if text:
                parts.append(_apply_fmt(text, fmt))

        elif local == "ins":
            # Tracked insertion — include text with its formatting
            for r in child.findall(f"{WNS}r"):
                fmt = _rpr_fmt(r.find(f"{WNS}rPr"))
                text = _run_text(r)
                if text:
                    parts.append(_apply_fmt(text, fmt))

        elif local == "del":
            # Tracked deletion — render as strikethrough
            for r in child.findall(f"{WNS}r"):
                dt = r.find(f"{WNS}delText")
                if dt is not None and dt.text:
                    parts.append(f"~~{dt.text}~~")

        elif local == "hyperlink":
            r_id = child.get(f"{RNS}id")
            url = None
            if r_id:
                rel = doc.part.rels.get(r_id)
                if rel and rel.is_external:
                    url = rel.target_ref
            link_parts = []
            for r in child.findall(f".//{WNS}r"):
                fmt = _rpr_fmt(r.find(f"{WNS}rPr"))
                text = _run_text(r)
                if text:
                    link_parts.append(_apply_fmt(text, fmt))
            link_text = "".join(link_parts)
            if link_text:
                parts.append(f"[{link_text}]({url})" if url else link_text)

    return "".join(parts)


def _table_to_md(table, doc):
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text_parts = []
            for para in cell.paragraphs:
                t = _para_to_md_text(para._element, doc).strip()
                if t:
                    cell_text_parts.append(t)
            cells.append(" ".join(cell_text_parts).replace("|", "\\|"))
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
    return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])


def _extract_paragraph_images(para, doc, images_dir, img_counter, img_rel_path=""):
    refs = []
    for run in para.runs:
        for drawing in run._element.findall(f".//{WNS}drawing"):
            for blip in drawing.iter(f"{DNS}blip"):
                embed = blip.get(f"{RNS}embed")
                if not embed:
                    continue
                rel = doc.part.rels.get(embed)
                if rel is None:
                    continue
                ext_map = {
                    "image/png": ".png", "image/jpeg": ".jpg",
                    "image/gif": ".gif", "image/bmp": ".bmp",
                    "image/tiff": ".tiff", "image/svg+xml": ".svg",
                }
                ext = ext_map.get(rel.target_part.content_type, ".png")
                img_name = f"image_{img_counter[0]:03d}{ext}"
                img_counter[0] += 1
                (images_dir / img_name).write_bytes(rel.target_part.blob)
                refs.append(f"![{img_name}]({img_rel_path}/{img_name})")
    return refs


def docx_to_md(docx_path, images_dir=None, img_rel_path=""):
    doc = Document(docx_path)
    lines = []
    img_counter = [1]

    table_elements = {tbl._element: tbl for tbl in doc.tables}
    table_inserted = set()

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]

        if tag == "p":
            para = next((p for p in doc.paragraphs if p._element is block), None)
            if para is None:
                continue

            if images_dir:
                for ref in _extract_paragraph_images(para, doc, images_dir, img_counter, img_rel_path):
                    lines.append(ref)

            text = _para_to_md_text(para._element, doc).strip()
            if not text:
                continue

            level = _heading_level(para)
            if level:
                lines.append(f"\n{'#' * level} {text}\n")
            elif _is_list_item(para):
                lines.append(f"1. {text}" if _is_ordered_list(para) else f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            if block in table_elements and id(block) not in table_inserted:
                table_inserted.add(id(block))
                lines.append("")
                lines.append(_table_to_md(table_elements[block], doc))
                lines.append("")

    return "\n".join(lines)


def main():
    sys.stdout.reconfigure(encoding="utf-8")

    parser = argparse.ArgumentParser(description="Convert .docx to markdown")
    parser.add_argument("input", help="Path to .docx file")
    parser.add_argument("-o", "--output", help="Output path (default: same name .md)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"File not found: {input_path}")
        sys.exit(1)

    output_path = Path(args.output) if args.output else input_path.with_suffix(".md")
    md = docx_to_md(input_path)
    output_path.write_text(md, encoding="utf-8")
    print(f"Saved: {output_path} ({len(md)} chars)")


if __name__ == "__main__":
    main()
